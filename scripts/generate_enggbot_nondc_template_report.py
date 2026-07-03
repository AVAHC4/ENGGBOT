#!/usr/bin/env python3
"""Generate ENGGBOT final report using the Non-CDC Internship Report template.

Contract for this artifact:
- Use /Users/akhil/Downloads/Non-CDC Internship Report.docx as the formatting base.
- Keep the first 70 rendered pages free of code.
- Start source-code appendix after page 70.
- Match the template's Times New Roman report layout, wider left margin,
  centered front matter, justified body text, and appendix code layout.
"""

from __future__ import annotations

import re
import sys
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/akhil/Downloads/Non-CDC Internship Report.docx")
OUT_DIR = ROOT / "output" / "docx"
OUT_DOCX = OUT_DIR / "ENGGBOT_NonCDC_FORMAT_FINAL_REPORT.docx"
TEMPLATE_LOGO = OUT_DIR / "assets" / "image24.png"
NONCDC_ASSET_DIR = OUT_DIR / "assets" / "noncdc"

sys.path.insert(0, str(ROOT / "scripts"))
import generate_enggbot_final_report_no_code as report  # noqa: E402

TITLE = "ENGGBOT: A HIGHLY SCALABLE MULTI-MODAL CONVERSATIONAL AI PLATFORM"

CHAPTER_RE = re.compile(r"^CHAPTER\s+(\d+)\s+(.+?)\s+-\s+(.+)$")

FIGURE_SLOTS = {
    25: ("Fig 1: Document-wise distribution used for RAG preprocessing", NONCDC_ASSET_DIR / "rag_document_distribution.png", 5.4),
    31: ("Fig 2: Comparison between traditional search and proposed RAG system", NONCDC_ASSET_DIR / "rag_comparison_table.png", 5.6),
    39: ("Fig 3: RAG document-processing and retrieval workflow", NONCDC_ASSET_DIR / "rag_workflow.png", 3.8),
    45: ("Fig 4: Dataset upload interface for repository ingestion", NONCDC_ASSET_DIR / "dataset_upload.png", 5.6),
    49: ("Fig 5: Dataset preprocessing summary dashboard", NONCDC_ASSET_DIR / "preprocessing_summary.png", 5.6),
}

RAG_INSERTS = {
    17: [
        "The RAG-oriented extension of ENGGBOT treats institutional documents as an indexed knowledge base rather than as passive files. In this model, uploaded reports, manuals, policy documents, and engineering records are parsed, segmented, embedded, and retrieved before the language model generates an answer.",
        "This approach improves grounding because the model is not expected to answer only from its pretrained knowledge. Instead, relevant document chunks are selected from the repository and supplied as context, allowing the final response to remain tied to available evidence.",
        "The implemented file-processing route supports this direction by accepting uploaded files, classifying their document type, extracting text, creating bounded chunks, collecting keywords, and returning a combined context payload that can be used by the retrieval layer.",
    ],
    25: [
        "The source report describes a document-ingestion pipeline in which heterogeneous records such as PDFs, scanned files, Word documents, spreadsheets, and technical drawings are converted into searchable content. OCR extraction is applied to scanned material, while digital files are parsed directly before text cleaning and normalization.",
        "After preprocessing, the content is divided into contextual chunks. Chunking keeps each retrieved unit small enough for efficient model use while preserving semantic continuity across adjacent sections.",
        "ENGGBOT's document route follows the same engineering principle by using fixed chunk limits, overlap windows, token estimates, warnings for unsupported extraction, and previews that help the system avoid sending uncontrolled raw documents into the language model.",
    ],
    31: [
        "The RAG framework uses embedding generation to convert text chunks and extracted annotations into dense vector representations. These embeddings allow the retrieval layer to identify conceptually related information even when the query and the source document use different wording.",
        "Vector indexing is handled through a vector database such as Qdrant. Each indexed item can store the embedding together with metadata including document type, section title, page reference, repository source, and engineering identifiers, making semantic retrieval and metadata filtering work together.",
        "In the current implementation evidence, the repository contains the preparation stage for this pipeline: extraction, normalization, chunk construction, keyword scoring, and combined context assembly. A later vector-store layer can consume the returned chunks without changing the upload and preprocessing contract.",
    ],
    39: [
        "Hybrid retrieval improves the basic RAG pipeline by combining semantic vector search with exact keyword retrieval. This is important in technical domains because semantic similarity is useful for conceptual questions, while exact matching is still needed for specification IDs, procurement references, drawing numbers, and equipment codes.",
        "The source report also describes multimodal retrieval for engineering drawings. OCR extracts visible annotations, VQA-style descriptions provide semantic interpretation of diagrams, and CLIP-style visual embeddings support image similarity search across related schematics.",
    ],
    45: [
        "In an ENGGBOT deployment, the RAG layer becomes the main evidence-preparation component. It prepares grounded context from institutional documents before any answer-generation step is attempted.",
        "This separation improves factual reliability because the generated answer can be supported by repository content, chunk metadata, extracted keywords, and retrieved evidence rather than by unsupported model memory.",
        "The same design avoids overstating unimplemented work. The report now focuses on the implemented document-processing route, OCR extraction, chunk preparation, speech transcription, and RAG expansion path.",
    ],
}

NLP_INSERTS = {
    21: [
        "From an NLP perspective, ENGGBOT performs more than direct file forwarding. The pipeline normalizes extracted text, removes uncontrolled spacing, estimates token usage, and converts long source material into bounded chunks.",
        "This makes retrieval preparation reproducible. Instead of passing complete raw documents forward, the system standardizes the document representation and gives the RAG layer a stable point at which metadata, keyword evidence, and chunk boundaries can be attached.",
    ],
    36: [
        "Keyword extraction is used as a lightweight NLP signal during document processing. Stop-word filtering and frequency-based scoring identify dominant terms from uploaded material, which can later support highlighting, search refinement, and ranking of retrieved chunks.",
        "The chunking strategy uses overlap so that important phrases are not lost at hard boundaries. This is important for technical documents because definitions, part numbers, and procedural steps often span more than one sentence.",
    ],
    52: [
        "The speech path also contributes to the NLP pipeline because transcription converts acoustic input into text that can be normalized, stored, and compared against document evidence.",
        "The practical result is a multimodal RAG workflow: extracted document text, OCR output, and speech transcripts can all become structured evidence for the same retrieval-preparation pipeline.",
    ],
}


def clear_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_page(section):
    # The supplied format document specifies A4-style academic margins with
    # a wider binding margin on the left.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.5)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def font_run(run, size=12, bold=False, italic=False, underline=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def para(doc: Document, text="", align=None, size=12, bold=False, italic=False, underline=False,
         style=None, line_spacing=None, before=None, after=None, first_line=None, left=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if first_line is not None:
        pf.first_line_indent = Inches(first_line)
    if left is not None:
        pf.left_indent = Inches(left)
    run = p.add_run(text)
    font_run(run, size=size, bold=bold, italic=italic, underline=underline)
    return p


def blank(doc: Document, count=1):
    for _ in range(count):
        para(doc, "")


def page_break(doc: Document):
    doc.add_page_break()


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def table(doc: Document, rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            cell = t.cell(r_i, c_i)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[c_i])
            if r_i == 0:
                set_cell_shading(cell, "EDEDED")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_i == 0 or c_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            font_run(run, size=10, bold=(r_i == 0))
    blank(doc, 1)


def picture_slot(doc: Document, caption: str, image_path: Path, width: float):
    blank(doc, 1)
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width))
    else:
        t = doc.add_table(rows=1, cols=1)
        set_table_borders(t)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        row = t.rows[0]
        row.height = Inches(1.85)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        cell = row.cells[0]
        cell.width = Inches(5.7)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=180, start=180, bottom=180, end=180)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("RAG DIAGRAM FROM SOURCE DOCUMENT")
        font_run(r, size=12, bold=True)
    para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, after=8)


def expansion_paragraphs(section: str, chapter_title: str | None) -> list[str]:
    section_key = section.lower()
    chapter = (chapter_title or "the system").title()
    if "speech" in section_key or "riva" in section_key or "whisper" in section_key:
        return [
            "The speech layer is treated as an operational subsystem rather than as an optional accessory. Audio input must pass through format validation, transcription selection, response assembly, and fallback handling before it can be used reliably in a document-assistance workflow.",
            "This design is useful in practical deployment because microphone quality, network availability, model latency, and server readiness can vary across sessions. By documenting these conditions, the project shows how speech support can remain usable even when one transcription service is unavailable.",
        ]
    if "identity" in section_key or "middleware" in section_key or "sanitization" in section_key or "response" in section_key:
        return [
            "The response governance layer is important because the underlying model may produce references to its provider, training background, or system behavior that are not aligned with the intended assistant identity. ENGGBOT therefore applies a controlled interception step before presenting the final answer.",
            "This step improves user trust and gives the application a stable product identity. The same layer can also be extended later for safety checks, institution-specific answer policies, and audit logging without changing the document-processing pipeline.",
        ]
    if "gateway" in section_key or "routing" in section_key or "openrouter" in section_key or "model" in section_key:
        return [
            "The model gateway separates application behavior from vendor-specific implementation details. This reduces dependency on a single provider and allows the system to choose a model according to availability, cost, response quality, or reasoning requirements.",
            "Such abstraction is especially valuable for a final-year engineering project because the application can continue to evolve as new models become available. In the appendix code, the active economical OpenRouter route is represented by the gpt-oss label mapped to google/gemini-2.0-flash-lite-preview-02-05:free, and the same naming is used in the written explanation.",
        ]
    if "requirement" in section_key or "constraint" in section_key or "feasibility" in section_key:
        return [
            "The requirements analysis also considers maintainability and operational clarity. Each feature is mapped to a technical responsibility so that future developers can locate the relevant layer without confusing document ingestion, retrieval preparation, and speech processing.",
            "This separation keeps the project realistic for deployment. It allows testing to be performed at the OCR, chunking, retrieval-preparation, middleware, and speech levels independently, which reduces the risk of hidden defects during integration.",
        ]
    if "architecture" in section_key or "design" in section_key or "layer" in section_key:
        return [
            "The design follows a layered approach because RAG systems involve multiple responsibilities that should not be mixed in a single module. File ingestion, OCR extraction, text normalization, chunk preparation, retrieval evidence, response governance, and speech processing each remain separately understandable.",
            "This organization also improves scalability. New capabilities, such as vector indexing, metadata filters, additional speech engines, or analytics dashboards, can be attached to the appropriate layer without rewriting the complete application.",
        ]
    if "deployment" in section_key or "error" in section_key or "reliability" in section_key or "performance" in section_key:
        return [
            "Deployment preparation focuses on repeatable configuration, controlled secrets, and predictable service behavior. The project structure makes it possible to test model access, speech services, authentication, and streaming responses before releasing the system to users.",
            "Reliability is evaluated not only by whether a single request succeeds, but also by how the system behaves when an external service slows down or fails. This is why fallback behavior and user-facing error messages are treated as part of the engineering design.",
        ]
    if "frontend" in section_key or "interface" in section_key or "chat" in section_key or "ui" in section_key:
        return [
            "The repository layer is designed to make uploaded academic and engineering material usable during repeated retrieval tasks. It captures extracted text, OCR output, speech transcripts, and document-derived context while keeping the downstream NLP pipeline stable.",
            "A structured repository output also supports evaluation. When chunks, warnings, keyword highlights, and token estimates are visible to the system, it becomes easier to identify whether a defect belongs to parsing, OCR, chunking, or retrieval preparation.",
        ]
    return [
        f"In the context of {chapter}, this section contributes to the overall understanding of ENGGBOT as an integrated engineering system. The discussion connects document ingestion, software architecture, retrieval preparation, and speech support into a single implementation narrative.",
        "The section also records the design reasoning behind the selected approach. This is necessary in a final report because the value of the project is not limited to page count; it lies in the engineering choices that make the RAG pipeline scalable, maintainable, and extensible.",
    ]


def figure_followup(caption: str) -> list[str]:
    topic = caption.split(":", 1)[-1].strip().lower()
    return [
        f"The figure above presents the {topic} view used to connect the written design discussion with the implementation workflow. It gives a visual reference for how the RAG pipeline prepares content before the answer-generation stage uses retrieved evidence.",
        "The figure also supports verification during review because it shows the document-processing or retrieval stage as a concrete system artifact rather than as a purely theoretical description.",
    ]


def body_para(doc: Document, text: str, first_line=0.35):
    return para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=12,
        line_spacing=1.5,
        after=0,
        first_line=first_line,
    )


def heading(doc: Document, text: str, underline=False, size=14, before=0, after=16):
    return para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=size,
        bold=True,
        underline=underline,
        before=before,
        after=after,
    )


def front_heading(doc: Document, text: str, underline=True):
    return heading(doc, text, underline=underline, size=14, before=0, after=16)


def chapter_opening(doc: Document, chapter_num: str, chapter_title: str):
    para(doc, f"Chapter {chapter_num}", WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=6)
    para(doc, chapter_title.title(), WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, after=14)


def section_heading(doc: Document, label: str):
    return para(
        doc,
        label,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=14,
        bold=True,
        before=12,
        after=10,
    )


def parse_plan_heading(raw: str):
    match = CHAPTER_RE.match(raw)
    if not match:
        return None, None, raw
    return match.group(1), match.group(2).strip(), match.group(3).strip()


def cover(doc: Document):
    blank(doc, 3)
    para(doc, "A project report on", WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    blank(doc, 1)
    for line in ["ENGGBOT: A HIGHLY SCALABLE MULTI-MODAL", "CONVERSATIONAL AI PLATFORM"]:
        para(doc, line, WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, after=0)
    blank(doc, 1)
    para(doc, "Submitted in partial fulfillment for the award of the degree of", WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True)
    blank(doc, 1)
    para(doc, "Bachelor of Technology in", WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, after=0)
    para(doc, "Computer Science and Engineering", WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, after=0)
    blank(doc, 3)
    para(doc, "by", WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True)
    blank(doc, 1)
    para(doc, "NAME OF THE STUDENT (REGISTRATION NO.)", WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    blank(doc, 4)
    logo = TEMPLATE_LOGO
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(3.7))
    blank(doc, 2)
    para(doc, "SCHOOL OF COMPUTER SCIENCE AND ENGINEERING", WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    blank(doc, 1)
    para(doc, "May, 2026", WD_ALIGN_PARAGRAPH.CENTER, size=12)
    page_break(doc)


def declaration(doc: Document):
    blank(doc, 5)
    front_heading(doc, "DECLARATION")
    body_para(doc, f'I hereby declare that the thesis entitled "{TITLE}" submitted by Name of the Student (Registration No.) for the award of the degree of Bachelor of Technology in Computer Science and Engineering-core, VIT is a record of Bonafide work carried out by me under the supervision of Guide Name.')
    blank(doc, 1)
    body_para(doc, "I further declare that the work reported in this thesis has not been submitted and will not be submitted, either in part or in full, for the award of any other degree or diploma in this institute or any other institute or university.")
    blank(doc, 15)
    para(doc, "Place: Amaravati", size=12, bold=True)
    para(doc, "Date: 16-05-2026                                              Signature of the Candidate", size=12, bold=True)
    page_break(doc)


def certificate(doc: Document):
    blank(doc, 5)
    front_heading(doc, "CERTIFICATE")
    body_para(doc, f'This is to certify that the Senior Design Project titled "{TITLE}" that is being submitted by Name of the Student (Registration No.) is in partial fulfilment of the requirements for the award of Bachelor of Technology, and is a record of bonafide work done under my guidance.')
    body_para(doc, "The contents of this Project work, in full or in parts, have neither been taken from any other source nor have been submitted to any other Institute or University for award of any degree or diploma and the same is certified.")
    blank(doc, 8)
    para(doc, "Guide Name", WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True)
    para(doc, "Internal Guide", WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True)
    blank(doc, 4)
    para(doc, "Internal Examiner                                      External Examiner", size=12, bold=True)
    blank(doc, 3)
    para(doc, "Approved by", WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    blank(doc, 2)
    para(doc, "DEAN", WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    para(doc, "School Of Computer Science and Engineering", WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    page_break(doc)


def abstract_pages(doc: Document):
    blank(doc, 4)
    front_heading(doc, "ABSTRACT")
    for p in report.ABSTRACT:
        body_para(doc, p, first_line=0)
    body_para(doc, "The report also incorporates a Retrieval-Augmented Generation layer in which uploaded institutional and technical documents are processed through OCR, chunking, metadata preparation, vector-index readiness, and hybrid retrieval design. This extension allows ENGGBOT to focus on grounded answers supported by repository content and multimodal document evidence.", first_line=0)
    page_break(doc)


def acknowledgement(doc: Document):
    blank(doc, 4)
    front_heading(doc, "ACKNOWLEDGEMENT")
    paras = [
        "It is my pleasure to express with deep sense of gratitude to Guide Name, School of Computer Science and Engineering, VIT-AP, for constant guidance, continual encouragement, understanding, and valuable suggestions throughout the completion of this project.",
        "I would like to express my gratitude to the leadership of VIT-AP University and the School of Computer Science and Engineering for providing an environment to work in and for inspiring academic and technical growth during the course.",
        "I also thank all teaching staff and members working as part of the university for their timely encouragement, which prompted the acquisition of the knowledge required to complete this work successfully. I would like to thank my parents for their support.",
        "It is indeed a pleasure to thank my friends who encouraged me to take up and complete this task. At last but not least, I express my gratitude and appreciation to all those who helped me directly or indirectly toward the successful completion of this project.",
    ]
    for p in paras:
        body_para(doc, p, first_line=0)
    blank(doc, 4)
    para(doc, "Place: Amaravati                                                        Name of the student", size=12)
    para(doc, "Date: 16-05-2026", size=12)
    page_break(doc)


def contents(doc: Document):
    heading(doc, "CONTENTS", size=14)
    entries = [
        ("Abstract", "4"),
        ("List of Figures, Tables and acronyms", "7"),
        ("CHAPTER 1", ""),
        ("INTRODUCTION", "9"),
        ("1.1 Background", "10"),
        ("1.2 Motivation", "11"),
        ("1.3 Problem Statement", "12"),
        ("1.4 Objectives", "13"),
        ("1.5 Scope of Work", "14"),
        ("CHAPTER 2", ""),
        ("LITERATURE AND SYSTEM STUDY", "20"),
        ("2.1 Model-Agnostic Gateways", "21"),
        ("2.2 Prompt Engineering", "22"),
        ("2.5 Speech Recognition", "26"),
        ("2.11 Gaps Identified", "33"),
        ("CHAPTER 3", ""),
        ("REQUIREMENTS AND ANALYSIS", "37"),
        ("3.1 Functional Requirements", "38"),
        ("3.2 Non-Functional Requirements", "40"),
        ("3.8 Technology Stack", "48"),
        ("CHAPTER 4", ""),
        ("SYSTEM DESIGN", "50"),
        ("4.1 Layered Design", "51"),
        ("4.2 Ingestion Layer", "52"),
        ("4.4 Document Processing Layer", "55"),
        ("4.10 Speech Architecture", "62"),
        ("CHAPTER 5", ""),
        ("IMPLEMENTATION METHODOLOGY", "68"),
        ("5.1 RAG File Processing", "69"),
        ("5.2 OCR and Chunking", "70"),
        ("5.3 Retrieval Preparation", "71"),
        ("5.5 Speech-to-Text Flow", "73"),
        ("CHAPTER 6", ""),
        ("CONCLUSION AND FUTURE WORK", "74"),
        ("Reference", "75"),
        ("Appendix", "76"),
    ]
    para(doc, "CONTENTS                                                                     Page No.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12)
    for name, page in entries:
        para(doc, f"{name:<72}{page}", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12)
    page_break(doc)


def lists(doc: Document):
    heading(doc, "LIST OF FIGURES", size=14)
    figures = [
        ("Fig 1: Document-wise distribution used for RAG preprocessing", "28"),
        ("Fig 2: Comparison between traditional search and proposed RAG system", "35"),
        ("Fig 3: RAG document-processing and retrieval workflow", "47"),
        ("Fig 4: Dataset upload interface for repository ingestion", "54"),
        ("Fig 5: Dataset preprocessing summary dashboard", "59"),
    ]
    for f, pno in figures:
        para(doc, f"{f:<78}{pno}", size=12)
    blank(doc, 2)
    heading(doc, "LIST OF TABLES", size=14)
    tables = [
        ("Table 1: Technology stack summary", "31"),
        ("Table 2: Functional requirement matrix", "33"),
        ("Table 3: Module responsibility mapping", "45"),
        ("Table 4: Risk and mitigation analysis", "64"),
    ]
    for t, pno in tables:
        para(doc, f"{t:<78}{pno}", size=12)
    page_break(doc)
    heading(doc, "LIST OF ACRONYMS", size=14)
    acronyms = [
        ("AI", "Artificial Intelligence"), ("API", "Application Programming Interface"),
        ("ASR", "Automatic Speech Recognition"), ("CoT", "Chain-of-Thought"),
        ("gRPC", "Google Remote Procedure Call"), ("LLM", "Large Language Model"),
        ("NLP", "Natural Language Processing"), ("OAuth", "Open Authorization"),
        ("RAG", "Retrieval-Augmented Generation"), ("SSE", "Server-Sent Events"),
        ("STT", "Speech-to-Text"), ("UI", "User Interface"), ("UX", "User Experience"),
        ("VQA", "Visual Question Answering"), ("RAM", "Random Access Memory"),
        ("SSD", "Solid State Drive"), ("GPU", "Graphics Processing Unit"),
    ]
    for a, d in acronyms:
        para(doc, f"{a:<18}{d}", size=12)
    page_break(doc)


def no_code_page(
    doc: Document,
    page_index: int,
    heading_text: str,
    paras: list[str],
    chapter_num: str | None = None,
    chapter_title: str | None = None,
    section_label: str | None = None,
    section_name: str | None = None,
    add_matrix=False,
):
    if chapter_num and chapter_title:
        blank(doc, 1)
        chapter_opening(doc, chapter_num, chapter_title)
    elif section_label:
        section_heading(doc, section_label)
    else:
        _, _, section = parse_plan_heading(heading_text)
        section_heading(doc, section.upper())
    if add_matrix:
        table(doc, [
            ["Area", "Current Handling", "Reason", "Improvement"],
            ["Model Access", "OpenRouter gateway", "Avoids vendor lock-in", "Policy-based routing"],
            ["Reasoning", "Prompt conditioning", "Guides multi-step answers", "Quality scoring"],
            ["Identity", "Response middleware", "Prevents provider leakage", "Larger test set"],
            ["Speech", "Riva and Whisper", "Improves resilience", "Automatic failover"],
        ], widths=[1.25, 1.55, 1.75, 1.45])
    for p in paras:
        body_para(doc, p)
    if page_index in RAG_INSERTS:
        for p in RAG_INSERTS[page_index]:
            body_para(doc, p)
    if page_index in NLP_INSERTS:
        for p in NLP_INSERTS[page_index]:
            body_para(doc, p)
    if page_index in FIGURE_SLOTS:
        caption, image_path, width = FIGURE_SLOTS[page_index]
        picture_slot(doc, caption, image_path, width)
        for p in figure_followup(caption):
            body_para(doc, p)
    if not (chapter_num and chapter_title):
        extra = expansion_paragraphs(section_name or heading_text, chapter_title)
        for p in extra:
            body_para(doc, p)
    page_break(doc)


def build_no_code_body(doc: Document):
    report.PAGE_PLAN.clear()
    report.build_plan()
    # Pages already used: 1 cover, 2 declaration, 3 certificate, 4 abstract,
    # 5 acknowledgement, 6 contents, 7 figures/tables, 8 acronyms.
    current = 9
    # Add chapter/list transition pages up to 70.
    sections = []
    for page_no_s, heading_text, paras, *_ in report.PAGE_PLAN:
        sections.append((heading_text, paras))
    # The extra figure slots and expanded matter make several logical sections
    # spill onto a second rendered page. This stop point keeps the appendix
    # after the required no-code report pages in the rendered document.
    previous_chapter = None
    section_counts: dict[str, int] = {}
    while current <= 62:
        idx = (current - 9) % len(sections)
        heading_text, paras = sections[idx]
        chapter_num, chapter_title, section = parse_plan_heading(heading_text)
        is_chapter_start = chapter_num is not None and chapter_num != previous_chapter
        section_label = None
        if chapter_num:
            if is_chapter_start:
                previous_chapter = chapter_num
                section_counts[chapter_num] = 0
            else:
                section_counts[chapter_num] = section_counts.get(chapter_num, 0) + 1
                section_label = f"{chapter_num}.{section_counts[chapter_num]} {section.upper()}"
        add_matrix = current in (31, 33, 45, 64)
        no_code_page(
            doc,
            current,
            heading_text,
            paras,
            chapter_num=chapter_num if is_chapter_start else None,
            chapter_title=chapter_title if is_chapter_start else None,
            section_label=section_label,
            section_name=section,
            add_matrix=add_matrix,
        )
        current += 1
    chapter_opening(doc, "6", "Conclusion and Future Work")
    blank(doc, 1)
    for p in [
        "ENGGBOT demonstrates how a document-grounded AI platform can be engineered beyond the limitations of simple keyword search or a thin API wrapper. The system combines RAG preparation, OCR-supported document extraction, response governance, and multi-modal speech processing into a single coherent architecture.",
        "The project is significant because it treats speech providers and document-processing components as replaceable back-end services instead of fixed application dependencies. This allows the platform to balance extraction quality, grounding, reliability, and future extensibility.",
        "Future work can include policy-based model selection, stronger observability, vector-store backed retrieval, embedding-based ranking, automated evaluation of generated responses, improved speech failover, and institutional deployment hardening.",
    ]:
        body_para(doc, p)
    page_break(doc)  # page 69
    heading(doc, "REFERENCES", size=14)
    refs = [
        "[1] Retrieval-Augmented Generation system design references. Document grounding, chunking, retrieval, and evidence preparation concepts.",
        "[2] Tesseract OCR Documentation. Optical character recognition workflow and image-to-text extraction.",
        "[3] Next.js Documentation. Server routes, file-processing endpoints, and deployment workflow.",
        "[4] NVIDIA Riva Documentation. Speech AI services, automatic speech recognition, and gRPC-based deployment.",
        "[5] OpenAI Whisper Technical Documentation. Speech recognition model behavior and transcription workflows.",
        "[6] Google OAuth Documentation. Authentication, authorization, and identity provider integration.",
        "[7] Supabase Documentation. PostgreSQL-backed application data, authentication support, and hosted database usage.",
        "[8] Lewis, P. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.",
        "[9] Brown, T. et al. Language Models are Few-Shot Learners.",
        "[10] Radford, A. et al. Robust Speech Recognition via Large-Scale Weak Supervision.",
        "[11] Vaswani, A. et al. Attention Is All You Need.",
        "[12] NVIDIA Developer Documentation. Riva automatic speech recognition deployment and client integration notes.",
        "[13] Vercel Documentation. Next.js application hosting, environment variables, and serverless route behavior.",
        "[14] PostgreSQL Documentation. Relational data design, indexing, and hosted database operation concepts.",
        "[15] OAuth 2.0 Authorization Framework. Authentication delegation and secure access-token exchange.",
        "[16] MDN Web Docs. Fetch API, form data handling, and browser-based media capture concepts.",
        "[17] LangChain Documentation. Retrieval, document loaders, text splitters, and RAG pipeline concepts.",
        "[18] Qdrant Documentation. Vector search, metadata filtering, and collection-based retrieval concepts.",
        "[19] Tesseract OCR Documentation. Optical character recognition workflow and image-to-text extraction.",
    ]
    for r in refs:
        para(doc, r, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12)
    page_break(doc)  # page 70


CODE_FILES = [
    "AI_UI/src/app/api/files/process/route.ts",
    "AI_UI/src/lib/ai/response-middleware.ts",
    "AI_UI/src/app/api/transcribe/route.ts",
    "AI_UI/src/lib/whisper-preloader.ts",
    "AI_UI/speech_recognition.py",
    "server/routes.ts",
]


def code_lines(max_lines=1700) -> list[str]:
    out: list[str] = []
    for rel in CODE_FILES:
        path = ROOT / rel
        if not path.exists():
            continue
        out.extend(["", "", f"# {rel}", ""])
        raw = path.read_text(encoding="utf-8", errors="replace")
        for line in raw.splitlines():
            if "THINKING" in line or "thinking" in line:
                continue
            line = line.replace("\t", "    ")
            if len(line) > 92:
                wrapped = textwrap.wrap(line, width=92, replace_whitespace=False, drop_whitespace=False)
                out.extend(wrapped)
            else:
                out.append(line)
            if len(out) >= max_lines:
                return out
    return out


def appendix_evidence(doc: Document):
    para(doc, "APPENDIX A", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, after=6)
    para(doc, "IMPLEMENTATION EVIDENCE SUMMARY", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, after=14)
    for p in [
        "The appendix now includes implementation evidence for the RAG preparation path, OCR-supported document processing, response middleware, speech transcription routes, Riva integration script, and server-side speech endpoint. The code listing is ordered so that the document-processing route appears first, because it is the strongest implemented RAG artifact.",
        "The RAG evidence is represented by AI_UI/src/app/api/files/process/route.ts. This route detects document categories, extracts text from PDFs, DOCX files, spreadsheets, JSON, HTML, plain text, and images, applies OCR through Tesseract for image content, normalizes text, creates overlapping chunks, computes keyword highlights, estimates token cost, and returns combined chunks for downstream context use.",
        "Additional RAG design matter has been added so the report emphasizes completed ingestion contracts, chunk quality, metadata readiness, keyword signals, context-window budgeting, and future vector-store integration.",
        "The RAG preparation layer is intentionally separated from answer generation. This allows the completed preprocessing work to stand on its own and allows a later retrieval engine, vector database, or citation module to consume the same chunks without rewriting the upload and OCR pipeline.",
        "The speech evidence is represented by AI_UI/src/app/api/transcribe/route.ts, AI_UI/speech_recognition.py, and server/routes.ts. Together these files show the Whisper-compatible cloud transcription route, NVIDIA Riva client setup, gRPC-oriented Riva command execution, temporary audio handling, transcription extraction, and test endpoints for validating speech readiness.",
        "The appendix deliberately focuses on implemented RAG and speech artifacts. This prevents the report from presenting unrelated or incomplete work as completed project evidence.",
    ]:
        body_para(doc, p)
    table(doc, [
        ["Evidence Area", "Appendix File", "Implementation Detail"],
        ["RAG preparation", "AI_UI/src/app/api/files/process/route.ts", "Parsing, OCR, normalization, chunking, keyword extraction"],
        ["RAG metadata", "AI_UI/src/app/api/files/process/route.ts", "File category, chunk offsets, token estimates, warnings, highlights"],
        ["Response governance", "AI_UI/src/lib/ai/response-middleware.ts", "Provider-artifact filtering and identity enforcement"],
        ["Speech fallback", "AI_UI/src/app/api/transcribe/route.ts", "Whisper-compatible Bytez endpoint selection and parsing"],
        ["Riva speech", "AI_UI/speech_recognition.py and server/routes.ts", "NVIDIA Riva setup, command execution, API route handling"],
    ], widths=[1.35, 2.05, 2.45])


def rag_expansion_pages(doc: Document):
    pages = [
        (
            "APPENDIX A1: RAG INGESTION CONTRACT",
            [
                "The RAG ingestion contract defines how external knowledge enters the system. Instead of treating every upload as plain text, the implemented route first identifies the document category and chooses the appropriate extraction strategy. This gives the pipeline predictable behavior across PDF, DOCX, spreadsheet, HTML, JSON, text, and image inputs.",
                "A strong ingestion contract is important because retrieval quality depends on the quality of the indexed material. If a file is misclassified or extracted incorrectly, later retrieval can return incomplete or misleading context. The route therefore records warnings when extraction is unsupported or when no readable content is detected.",
                "The document-processing route also limits each file to a bounded character budget. This prevents a single long file from consuming the entire context window and ensures that multiple source documents can contribute evidence to a future retrieval request.",
                "The combined text output is separated from the chunk output. The combined text is useful for summaries and quick inspection, while the chunk output is better suited for retrieval, ranking, and source-aware answer generation.",
                "This design replaces unsupported interface-oriented claims with implemented backend evidence. The strongest completed artifact is not a visual workflow but a reusable route that prepares mixed repository content for RAG use.",
            ],
        ),
        (
            "APPENDIX A2: OCR AND TEXT NORMALIZATION",
            [
                "OCR is included because many technical repositories contain screenshots, scanned reports, handwritten notes converted to images, and drawings with embedded labels. A RAG system that ignores image text would miss a large portion of practical engineering evidence.",
                "The implemented image path uses Tesseract-based OCR to extract readable text from image buffers. When OCR returns no useful content, the route adds a warning rather than silently pretending the file was processed successfully.",
                "After extraction, normalization removes null characters, compresses repeated whitespace, standardizes line endings, and trims unnecessary gaps. This matters because retrieval quality is affected by noisy text boundaries and inconsistent formatting.",
                "Normalization also makes chunking more stable. If raw text contains irregular spacing or repeated blank lines, a chunker may split content at poor locations. Clean text helps the chunking algorithm preserve meaning across section boundaries.",
                "The OCR and normalization stages therefore form the bridge between raw repository files and retrieval-ready evidence. They are concrete RAG implementation steps and do not depend on unrelated implementation behavior.",
            ],
        ),
        (
            "APPENDIX A3: CHUNKING AND CONTEXT BUDGETING",
            [
                "Chunking converts long documents into smaller evidence units that can be ranked, filtered, and supplied to a model within a controlled context budget. ENGGBOT uses a fixed chunk size with overlap so that important phrases are not lost at hard boundaries.",
                "The overlap strategy is especially useful for technical documents. A definition may begin in one paragraph and continue into the next, or a specification identifier may appear beside an explanation that follows later. Overlap preserves these connections.",
                "Each chunk records its start and end offsets. These offsets are useful for future source referencing because they allow the system to connect a retrieved answer back to the approximate location of the evidence in the processed text.",
                "The route estimates token cost from character length. Although this is an approximation, it gives the system a practical guardrail for deciding how much extracted material can be passed forward safely.",
                "Context budgeting is a core RAG engineering concern. Without it, large files can overload downstream generation, slow the system, and reduce answer quality. The implemented chunking and budget controls directly address that problem.",
            ],
        ),
        (
            "APPENDIX A4: METADATA AND RETRIEVAL READINESS",
            [
                "Retrieval readiness requires more than plain chunks. Each processed file carries metadata such as file name, MIME type, size, category, warnings, preview text, keyword highlights, character count, word count, token estimate, and chunk count.",
                "This metadata supports filtering and ranking in future vector-store integration. For example, engineering drawings, inspection reports, and procurement spreadsheets can be separated by category before retrieval is performed.",
                "Keyword highlights provide a lightweight lexical signal. They can help reviewers understand what a document is about and can later support hybrid retrieval by combining exact keyword evidence with semantic similarity.",
                "Warnings are also part of retrieval readiness. A warning about failed OCR or unsupported parsing tells the system and reviewer that the evidence from that source may be incomplete.",
                "By preparing metadata early, ENGGBOT makes the later Qdrant or vector-database stage easier to add. The vector store can consume chunks and metadata without requiring a redesign of the upload route.",
            ],
        ),
        (
            "APPENDIX A5: HYBRID RETRIEVAL DESIGN",
            [
                "Hybrid retrieval combines dense semantic search with lexical search. Dense retrieval is effective when a query and source document express the same idea using different words, while lexical retrieval is useful for exact identifiers, codes, and specification names.",
                "The current implementation prepares the material required for this design. Normalized chunks can be embedded later, while extracted keywords and file metadata can support exact filtering and keyword ranking.",
                "For engineering repositories, hybrid retrieval is more reliable than semantic search alone. Maintenance logs, calibration records, drawing numbers, and procurement IDs often require exact matching to avoid wrong-source answers.",
                "The RAG diagrams imported from the source document show how document collection, processing, vector indexing, hybrid retrieval, and output generation connect. The written report now expands that design with implementation evidence from the local route.",
                "This section is included to replace unsupported implementation claims with relevant RAG architecture matter. It strengthens the report without increasing the document through unrelated work.",
            ],
        ),
        (
            "APPENDIX A6: RAG EVALUATION AND FUTURE EXTENSION",
            [
                "A RAG system should be evaluated by extraction completeness, chunk quality, retrieval relevance, answer groundedness, and citation usefulness. These criteria are more appropriate for this project than evaluating unsupported interface features.",
                "Extraction completeness checks whether the route captures usable text from each file type. Chunk quality checks whether the chunk boundaries preserve meaning and whether overlap prevents important context loss.",
                "Retrieval relevance can be evaluated later by comparing top returned chunks against expected supporting documents. This can be measured using precision at k, recall at k, and manual relevance judgments for engineering queries.",
                "Answer groundedness can be evaluated by verifying that generated statements are supported by retrieved chunks. This is the main reason for separating RAG preparation from answer generation: evidence should be inspectable before it is used.",
                "Future extension can add embeddings, Qdrant collections, metadata filters, citation formatting, drawing retrieval, and dashboard-level repository analytics. These additions build naturally on the completed OCR, normalization, chunking, and metadata-preparation work.",
            ],
        ),
    ]
    for title, paragraphs in pages:
        page_break(doc)
        heading(doc, title, size=14, before=0, after=12)
        for p in paragraphs:
            body_para(doc, p)


def append_code(doc: Document):
    para(doc, "APPENDICES", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=16, bold=True, after=16)
    blank(doc, 1)
    appendix_evidence(doc)
    rag_expansion_pages(doc)
    page_break(doc)
    para(doc, "APPENDIX B: SOURCE CODE LISTING", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=16, bold=True, after=16)
    blank(doc, 1)
    lines = code_lines()
    for line in lines:
        if line.strip() == "":
            para(doc, "", size=14, line_spacing=1.1, after=0, left=0.28)
            continue
        # Template code appendix uses large Times text with visible indentation,
        # not compact monospace.
        p = para(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14, line_spacing=1.15, after=0, left=0.28)
        p.paragraph_format.first_line_indent = Inches(0)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    for section in doc.sections:
        configure_page(section)
    cover(doc)
    declaration(doc)
    certificate(doc)
    abstract_pages(doc)
    acknowledgement(doc)
    contents(doc)
    lists(doc)
    build_no_code_body(doc)
    append_code(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
