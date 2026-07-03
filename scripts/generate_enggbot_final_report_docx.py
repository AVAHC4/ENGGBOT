#!/usr/bin/env python3
"""Generate the Word DOCX version of the ENGGBOT final report.

The content is sourced from generate_enggbot_final_report_no_code.py so the DOCX
matches the no-code PDF report in substance while using native Word structures.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "docx"
OUT_DOCX = OUT_DIR / "ENGGBOT_SDP_FINAL_REPORT.docx"

sys.path.insert(0, str(ROOT / "scripts"))
import generate_enggbot_final_report_no_code as report  # noqa: E402


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "333333"),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)

    footer_para = section.footer.paragraphs[0]
    add_page_number(footer_para)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, italic: bool = False,
                 space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_heading_page(doc: Document, heading: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(heading.upper())
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16 if len(heading) < 65 else 13)
    run.bold = True


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(9)
            run.bold = r_idx == 0
    doc.add_paragraph()


def page_break(doc: Document):
    doc.add_page_break()


def cover(doc: Document):
    add_centered(doc, "A project report on", 13, italic=True, space_after=28)
    for line in ["ENGGBOT: A HIGHLY SCALABLE MULTI-MODAL", "CONVERSATIONAL AI PLATFORM"]:
        add_centered(doc, line, 20, bold=True, space_after=12)
    add_centered(doc, "Submitted in partial fulfilment for the award of the degree of", 14, italic=True, space_after=28)
    add_centered(doc, "Bachelor Of Technology In", 22, bold=True, space_after=10)
    add_centered(doc, "Computer Science and Engineering", 22, bold=True, space_after=34)
    add_centered(doc, "by", 14, italic=True, space_after=24)
    add_centered(doc, "NAME OF THE STUDENT (REGISTRATION NO.)", 15, bold=True, space_after=16)
    logo = report.LOGO_WHITE
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(3.4))
    add_centered(doc, "SCHOOL OF COMPUTER", 16, bold=True, space_after=8)
    add_centered(doc, "SCIENCE AND ENGINEERING(SCOPE)", 16, bold=True, space_after=28)
    add_centered(doc, "May,2026", 12)
    page_break(doc)


def simple_page(doc: Document, heading: str, paragraphs: list[str], items: list[str] | None = None):
    add_heading_page(doc, heading)
    for text in paragraphs:
        add_body_paragraph(doc, text)
    if items:
        add_bullets(doc, items)
    page_break(doc)


def contents_pages(doc: Document):
    add_heading_page(doc, "CONTENTS")
    entries = [
        ("Abstract", "4-5"), ("List of Figures, Tables and Acronyms", "9-11"),
        ("CHAPTER 1 INTRODUCTION", "12"), ("CHAPTER 2 LITERATURE AND SYSTEM STUDY", "22"),
        ("CHAPTER 3 REQUIREMENTS AND ANALYSIS", "35"), ("CHAPTER 4 SYSTEM DESIGN", "45"),
        ("CHAPTER 5 IMPLEMENTATION METHODOLOGY", "60"), ("CHAPTER 6 RESULTS AND DISCUSSION", "70"),
        ("CHAPTER 7 CONCLUSION AND FUTURE WORK", "74"), ("References", "75"),
        ("Appendix: Non-Code Supporting Material", "76"),
    ]
    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "Contents"
    table.cell(0, 1).text = "Page No."
    for i, (name, page_no) in enumerate(entries, start=1):
        table.cell(i, 0).text = name
        table.cell(i, 1).text = page_no
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    page_break(doc)
    add_heading_page(doc, "LIST OF FIGURES, TABLES AND ACRONYMS")
    add_body_paragraph(doc, "The Word version keeps the same report content and major figure/table references as the final no-code PDF report.")
    add_bullets(doc, [
        "Figure references include system architecture, dual request-processing pipeline, UI verification, and module readiness.",
        "Table references include technology stack, functional requirements, non-functional requirements, module responsibilities, and risk analysis.",
        "Acronyms include AI, API, ASR, CoT, gRPC, LLM, NLP, OAuth, RAG, SSE, STT, UI, and UX.",
    ])
    page_break(doc)


def add_special_content(doc: Document, page_no: int, heading: str):
    if page_no in (37, 38, 41, 49, 52, 72):
        add_table(doc, [
            ["Area", "Requirement", "Current Handling", "Future Improvement"],
            ["Model Access", "Flexible provider support", "Gateway abstraction", "Add policy-based routing"],
            ["Reasoning", "Better structured answers", "Thinking Mode", "Evaluate answer quality"],
            ["Identity", "Consistent persona", "Middleware interception", "Expand test corpus"],
            ["Speech", "Reliable transcription", "Dual pipeline", "Add automatic provider selection"],
            ["Deployment", "Production readiness", "Build and env setup", "Add monitoring dashboard"],
        ])
    if page_no == 63 and (report.ASSET_DIR / "login.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(report.ASSET_DIR / "login.png"), width=Inches(5.1))
        cap = doc.add_paragraph("Fig 6: User interface verification screenshot")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def body_pages(doc: Document):
    report.PAGE_PLAN.clear()
    report.build_plan()
    for page_no_s, heading, paras, _fig, _table, _items in report.PAGE_PLAN:
        page_no = int(page_no_s)
        add_heading_page(doc, heading)
        add_special_content(doc, page_no, heading)
        for text in paras:
            add_body_paragraph(doc, text)
        page_break(doc)


def references_and_appendix(doc: Document):
    simple_page(doc, "REFERENCES", [
        "OpenRouter Documentation. Model routing, provider abstraction, and chat completion API concepts.",
        "React Documentation. Component-based user interface development and state-driven rendering.",
        "Next.js Documentation. Application routing, server routes, and deployment workflow.",
        "NVIDIA Riva Documentation. Speech AI services, automatic speech recognition, and gRPC-based deployment.",
        "OpenAI Whisper Technical Documentation. Speech recognition model behavior and transcription workflows.",
        "Google OAuth Documentation. Authentication, authorization, and identity provider integration.",
        "Supabase Documentation. PostgreSQL-backed application data, authentication support, and hosted database usage.",
        "Tailwind CSS Documentation. Utility-first styling and responsive interface development.",
        "Framer Motion Documentation. Motion primitives and animated React interfaces.",
        "Wei, J. et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models.",
        "Brown, T. et al. Language Models are Few-Shot Learners.",
        "Radford, A. et al. Robust Speech Recognition via Large-Scale Weak Supervision.",
    ])
    simple_page(doc, "APPENDIX: NON-CODE SUPPORTING MATERIAL", [
        "This appendix summarizes the non-code project evidence used in preparing the final report. The report is based on the implemented ENGGBOT repository, which includes the animated user interface, AI user interface, OpenRouter client, response middleware, speech recognition routes, authentication modules, project documentation, and deployment configuration.",
        "No source-code listing is included in this final report. The appendix is intentionally limited to descriptive supporting material so that the document remains a formal report rather than a code dump.",
        "Key project artifacts reviewed include the OpenRouter model gateway, Thinking Mode prompt behavior, middleware-based identity control, NVIDIA Riva speech support, Whisper fallback routes, chat context management, user interface components, and production build output.",
    ], items=[
        "Primary system focus: model-agnostic conversational AI platform.",
        "Major technical contribution: controlled orchestration of model routing, prompt behavior, identity enforcement, and speech fallback.",
        "Final report constraint: more than 70 pages of report material without code listings.",
    ])


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    report.prepare_logo()
    doc = Document()
    configure_doc(doc)
    cover(doc)
    simple_page(doc, "DECLARATION", [
        f'I hereby declare that the thesis entitled "{report.TITLE}" submitted by Name of the Student (Registration No.), for the award of the degree of Bachelor of Technology in Computer Science and Engineering, VIT-AP University, is a record of bonafide work carried out by me under the supervision of Guide Name.',
        "I further declare that the work reported in this thesis has not been submitted and will not be submitted, either in part or in full, for the award of any other degree or diploma in this institute or any other institute or university.",
        "Place: Amaravati\n\nDate:\n\nSignature of the Candidate",
    ])
    simple_page(doc, "CERTIFICATE", [
        f'This is to certify that the Senior Design Project titled "{report.TITLE}" submitted by Name of the Student (Registration No.) is in partial fulfilment of the requirements for the award of Bachelor of Technology and is a record of bonafide work done under my guidance.',
        "The contents of this Project work, in full or in parts, have neither been taken from any other source nor have been submitted to any other Institute or University for award of any degree or diploma and the same is certified.",
        "Guide Name\nInternal Guide\n\nThe thesis is satisfactory / unsatisfactory\n\nInternal Examiner                 External Examiner\n\nApproved by\n\nDEAN\nSchool Of Computer Science and Engineering",
    ])
    simple_page(doc, "ABSTRACT", report.ABSTRACT[:2])
    simple_page(doc, "ABSTRACT (CONTINUED)", report.ABSTRACT[2:])
    simple_page(doc, "ACKNOWLEDGEMENT", [
        "It is my pleasure to express a deep sense of gratitude to Guide Name, School of Computer Science and Engineering, VIT-AP University, for constant guidance, encouragement, and valuable suggestions throughout the completion of this project.",
        "I thank the faculty members, laboratory staff, classmates, and friends who supported the development and documentation of ENGGBOT. I also thank my parents for their constant support.",
        "Place: Amaravati\n\nDate:\n\nName of the student",
    ])
    contents_pages(doc)
    body_pages(doc)
    references_and_appendix(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
